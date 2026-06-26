"""
Format conversion utilities for download.

Provides conversion between markdown and other formats (txt, docx).
"""

import re
from io import BytesIO

from docx import Document
from docx.shared import Pt


def md_to_txt(md_content: str) -> str:
    """
    Convert markdown to plain text by stripping markdown syntax.

    Handles: headers, bold/italic, links, images, inline code, code blocks.
    """
    text = md_content

    # Remove headers (# ## ### etc.)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)

    # Remove bold/italic markers
    text = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', text)
    text = re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', text)

    # Remove links, keep text: [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

    # Remove images: ![alt](url) -> [alt] or empty
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'[\1]' if r'\1' else '', text)

    # Remove inline code backticks
    text = re.sub(r'`([^`]+)`', r'\1', text)

    # Remove code blocks (``` ... ```)
    text = re.sub(r'```[\s\S]*?```', '', text)

    # Remove horizontal rules
    text = re.sub(r'^[-*_]{3,}$', '', text, flags=re.MULTILINE)

    # Clean up excessive blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def md_to_docx(md_content: str, title: str = "Document") -> bytes:
    """
    Convert markdown to DOCX format.

    Handles: headers (h1-h6), bullet lists, numbered lists, paragraphs.
    """
    doc = Document()
    doc.core_properties.title = title

    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Handle headers (# to ######)
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            level = min(max(level, 1), 6)  # Clamp to 1-6
            text = stripped.lstrip('#').strip()
            if text:
                doc.add_heading(text, level=level)

        # Handle code blocks (skip content, just note it)
        elif stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_text = '\n'.join(code_lines)
                para = doc.add_paragraph()
                run = para.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

        # Handle bullet lists (- or *)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            doc.add_paragraph(text, style='List Bullet')

        # Handle numbered lists (1. 2. etc.)
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            doc.add_paragraph(text, style='List Number')

        # Handle blockquotes (>)
        elif stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            para = doc.add_paragraph(text)
            para.paragraph_format.left_indent = Pt(36)

        # Regular paragraph
        else:
            # Clean up inline markdown for paragraph text
            text = _clean_inline_markdown(stripped)
            if text:
                doc.add_paragraph(text)

        i += 1

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _clean_inline_markdown(text: str) -> str:
    """Remove inline markdown syntax from text."""
    # Remove bold/italic
    text = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', text)
    text = re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', text)
    # Remove links, keep text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Remove inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text
