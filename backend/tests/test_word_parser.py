from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw

from app.adapters.word import parse_word_document
from app.models.document_ir import SourceInfo
from app.pipeline.stages.normalize import NormalizeStage
from app.pipeline.stages.package import PackageStage


def _make_docx(path: Path, image_path: Path | None = None) -> None:
    document = Document()
    document.add_heading("台灣經濟研究院出差暨旅費報支辦法", level=1)
    document.add_paragraph("表四 台灣經濟研究院國內出差旅費報支數額表")
    table = document.add_table(rows=1, cols=5)
    headers = ["職稱/職級別", "交通費", "宿費平日", "宿費假日", "雜費"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    row = table.add_row().cells
    row[0].text = "院長、副院長、主任秘書"
    row[1].text = "按實檢據報支"
    row[2].text = "4500"
    row[3].text = "5500"
    row[4].text = "700"
    row = table.add_row().cells
    row[0].text = "研究員、副研究員、助理研究員"
    row[1].text = "按實檢據報支"
    row[2].text = "3500"
    row[3].text = "4500"
    row[4].text = "500"
    document.add_paragraph("圖一 台灣經濟研究院作業流程圖")
    if image_path:
        document.add_picture(str(image_path))
    document.add_paragraph("本表自114.1.1正式生效")
    document.save(path)


def test_parse_word_document_emits_structured_table(tmp_path):
    source = tmp_path / "travel.docx"
    _make_docx(source)

    result = parse_word_document(source, tmp_path / "out")

    assert result.success
    assert result.content_list_path is not None
    content = result.content_list_path.read_text(encoding="utf-8")
    assert "表四 台灣經濟研究院國內出差旅費報支數額表" in content
    assert "<th>職稱/職級別</th>" in content
    assert "<td>研究員、副研究員、助理研究員</td>" in content
    assert result.stats["parser"] == "native_docx"
    assert result.stats["table_count"] == 1


async def test_word_content_flows_to_package_markdown(tmp_path):
    source = tmp_path / "travel.docx"
    _make_docx(source)
    parse_result = parse_word_document(source, tmp_path / "out")
    assert parse_result.content_list_path is not None

    normalize_result = await NormalizeStage().run(
        doc_id="doc-docx",
        run_id="run-docx",
        content_list_path=parse_result.content_list_path,
        source_info=SourceInfo(
            path=str(source),
            ext=".docx",
            sha256="abc",
            size_bytes=source.stat().st_size,
        ),
        render_pages=False,
        mineru_version="native_docx",
    )
    assert normalize_result.success
    assert normalize_result.document_ir is not None

    package = PackageStage()
    source_md, _source_map = package._render_rag_md(
        document_ir=normalize_result.document_ir,
        asset_map={},
        enrichments={},
    )

    assert "表格名稱：表四 台灣經濟研究院國內出差旅費報支數額表" in source_md
    assert "欄位：職稱/職級別、交通費、宿費平日、宿費假日、雜費" in source_md
    assert "### 研究員、副研究員、助理研究員" in source_md
    assert "- 宿費平日：3500" in source_md
    assert "- 宿費假日：4500" in source_md
    assert "TABLE:" not in source_md
    assert "本表自114.1.1正式生效" in source_md



def _make_image(path: Path) -> None:
    image = Image.new("RGB", (240, 120), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((20, 30, 100, 80), outline="black", width=2)
    draw.rectangle((140, 30, 220, 80), outline="black", width=2)
    draw.line((100, 55, 140, 55), fill="black", width=2)
    image.save(path)


def test_parse_word_document_extracts_embedded_images(tmp_path):
    image_path = tmp_path / "flow.png"
    _make_image(image_path)
    source = tmp_path / "workflow.docx"
    _make_docx(source, image_path=image_path)

    result = parse_word_document(source, tmp_path / "out")

    assert result.success
    assert result.content_list_path is not None
    content = result.content_list_path.read_text(encoding="utf-8")
    assert '"type": "image"' in content
    assert '"img_path": "images/docx_img0000.png"' in content
    assert "圖一 台灣經濟研究院作業流程圖" in content
    assert (result.content_list_path.parent / "images" / "docx_img0000.png").exists()
    assert result.stats["image_count"] == 1


async def test_word_embedded_images_flow_to_document_ir(tmp_path):
    image_path = tmp_path / "flow.png"
    _make_image(image_path)
    source = tmp_path / "workflow.docx"
    _make_docx(source, image_path=image_path)
    parse_result = parse_word_document(source, tmp_path / "out")
    assert parse_result.content_list_path is not None

    normalize_result = await NormalizeStage().run(
        doc_id="doc-docx-image",
        run_id="run-docx-image",
        content_list_path=parse_result.content_list_path,
        source_info=SourceInfo(
            path=str(source),
            ext=".docx",
            sha256="abc",
            size_bytes=source.stat().st_size,
        ),
        render_pages=False,
        mineru_version="native_docx",
    )

    assert normalize_result.success
    assert normalize_result.document_ir is not None
    image_blocks = [
        block for block in normalize_result.document_ir.blocks
        if block.type.value == "image"
    ]
    assert len(image_blocks) == 1
    assert image_blocks[0].payload["img_path"] == "images/docx_img0000.png"
